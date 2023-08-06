import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from datetime import datetime

from docx import Document

from . import helper_funcs

def get_diagnostics_obj(source_shape, source_shape_id, target_shape, target_shape_id, intersect_unedited, intersect, tolerance_percent, tolerance_units, tolerance_values_sim_prop):
	units = helper_funcs.get_projection_units(source_shape, target_shape)

	source_smallest_area_intersect_id = intersect_unedited.sort_values('area_base_source')[source_shape_id].iloc[0]
	source_smallest_area_intersect = min(intersect_unedited['area_base_source'])
	source_total_area = sum(source_shape['area_base_source'])
	source_smallest_area_intersect_percent_total = source_smallest_area_intersect / source_total_area
	source_avg_area = np.mean(source_shape['area_base_source'])
	source_smallest_times_avg = source_avg_area / source_smallest_area_intersect

	target_smallest_area_intersect_id = intersect_unedited.sort_values('area_base_target')[target_shape_id].iloc[0]
	target_smallest_area_intersect = min(intersect_unedited['area_base_target'])
	target_total_area = sum(target_shape['area_base_target'])
	target_smallest_area_intersect_percent_total = target_smallest_area_intersect / target_total_area
	target_avg_area = np.mean(target_shape['area_base_target'])
	target_smallest_times_avg = target_avg_area / target_smallest_area_intersect

	intersect_smallest_area_from_source_and_target = min(min(intersect_unedited['area_base_source']), min(intersect_unedited['area_base_target']))
	intersect_smallest_area_id = intersect.sort_values('intersect_area')['INTERSECT_ID'].iloc[0]
	intersect_smallest_area = min(intersect['intersect_area'])
	intersect_total_area = sum(intersect['intersect_area'])
	intersect_smallest_area_percent_total = intersect_smallest_area / intersect_total_area
	intersect_avg_area = np.mean(intersect['intersect_area'])
	intersect_smallest_times_avg = intersect_avg_area / intersect_smallest_area

	tolerance_simulations_obj = get_tolerance_simulations_obj(tolerance_values_sim_prop, intersect_unedited)

	diagnostics_obj = {
		'title': 'Diagnostics',
		'units': units,
		'area_lost_during_intersection': {
			'area_lost_from_source': source_total_area - intersect_total_area,
			'area_lost_from_source_percent': (source_total_area - intersect_total_area)/source_total_area*100,
			'area_lost_from_target': target_total_area - intersect_total_area,
			'area_lost_from_target_percent': (target_total_area - intersect_total_area)/target_total_area*100
		},
		'source_shape': {
			'title': 'Statistics about source shape',
			'smallest_area': {
				'ID_name': source_shape_id,
				'ID_val': source_smallest_area_intersect_id,
				'area': source_smallest_area_intersect,
				'area_percent_total': source_smallest_area_intersect_percent_total,
			},
			'total_area': source_total_area,
			'avg_area': source_avg_area,
			'area_smallest_times_avg': source_smallest_times_avg,
			'plot': ''
		},
		'target_shape': {
			'title': 'Statistics about target shape',
			'smallest_area': {
				'ID_name': target_shape_id,
				'ID_val': target_smallest_area_intersect_id,
				'area': target_smallest_area_intersect,
				'area_percent_total': target_smallest_area_intersect_percent_total,
			},
			'total_area': target_total_area,
			'avg_area': target_avg_area,
			'area_smallest_times_avg': target_smallest_times_avg,
			'plot': ''
		},
		'intersect_shape': {
			'title': 'Statistics about intersected shape',
			# this note is a bit awkwardly placed here but it's okay for now
			'note': 'NB: The smallest area from either the source or the target shapes is {sa}. The smallest possible area in the intersect is {tp}% of that, or {tu} units.'.format(sa=helper_funcs.format_num(intersect_smallest_area_from_source_and_target), tp=helper_funcs.format_num(tolerance_percent, percent=True), tu=helper_funcs.format_num(tolerance_units)), 
			'smallest_area': {
				'ID_name': 'INTERSECT_ID',
				'ID_val': intersect_smallest_area_id,
				'area': intersect_smallest_area,
				'area_percent_total': intersect_smallest_area_percent_total,
			},
			'total_area': intersect_total_area,
			'avg_area': intersect_avg_area,
			'area_smallest_times_avg': intersect_smallest_times_avg,
			'plot': ''
		},
		'tolerance_simulations': {
			'title': 'Tolerance value simulations',
			'simulations': tolerance_simulations_obj,
			'plot': ''
		}
	}	

	# -----------------------------------------------------------plots
	# there is probably a better way to do this
	memfile_source_shape_plot = BytesIO()
	fig, ax = plt.subplots(1, 1)
	ax.set_title('Source shape')
	# set aspect to equal. This is done automatically
	# when using *geopandas* plot on it's own, but not when
	# working with pyplot directly.
	ax.set_aspect('equal')
	source_shape.plot(ax=ax)
	plt.savefig(memfile_source_shape_plot)
	diagnostics_obj['source_shape']['plot'] = memfile_source_shape_plot
	plt.close() # safely closing plt to prevent inline display of figures

	memfile_target_shape_plot = BytesIO()
	fig, ax = plt.subplots(1, 1)
	ax.set_title('Target shape')
	ax.set_aspect('equal')
	target_shape.plot(ax=ax)
	plt.savefig(memfile_target_shape_plot)
	diagnostics_obj['target_shape']['plot'] = memfile_target_shape_plot
	plt.close()

	memfile_intersect_shape_plot = BytesIO()
	fig, ax = plt.subplots(1, 1)
	if source_total_area >= target_total_area:
		ax.set_title('Source & Target shapes overlay\n(intersected area in brown and extra unused area from the source shape in pink)', fontsize=10, pad=20)
		ax.set_aspect('equal')
		source_shape.plot(ax=ax, zorder=2, color='red', alpha=0.6)
		target_shape.plot(ax=ax, zorder=1, color='green')
	elif source_total_area < target_total_area:
		ax.set_title('Source & Target shapes overlay\n(intersected area in brown and extra unused area from the target shape in pink)', fontsize=10, pad=20)
		ax.set_aspect('equal')
		target_shape.plot(ax=ax, zorder=2, color='red', alpha=0.6)
		source_shape.plot(ax=ax, zorder=1, color='green')
	plt.savefig(memfile_intersect_shape_plot)
	diagnostics_obj['intersect_shape']['plot'] = memfile_intersect_shape_plot
	plt.close()

	return diagnostics_obj

def get_diagnostics_obj_text(diagnostics_obj):
	diagnostics_obj_text = {
		'title': diagnostics_obj['title'],
		'units': f"The measurement unit for all reported numbers is: {diagnostics_obj['units']}",
		'area_lost_during_intersection': 'There were {area_lost_from_source} ({area_lost_from_source_percent}%) units area lost from the source shape, and {area_lost_from_target} ({area_lost_from_target_percent}%) units area lost from the target shape.'.format(area_lost_from_source=helper_funcs.format_num(diagnostics_obj['area_lost_during_intersection']['area_lost_from_source']), 
																																																																	area_lost_from_source_percent=helper_funcs.format_num(diagnostics_obj['area_lost_during_intersection']['area_lost_from_source_percent'], digits=2), 
																																																																	area_lost_from_target=helper_funcs.format_num(diagnostics_obj['area_lost_during_intersection']['area_lost_from_target']), 
																																																																	area_lost_from_target_percent=helper_funcs.format_num(diagnostics_obj['area_lost_during_intersection']['area_lost_from_target_percent'], digits=2)),
		'source_shape': {
			'title': diagnostics_obj['source_shape']['title'],
			'smallest_area': f"Smallest area (that is also found in the intersection): ID ({diagnostics_obj['source_shape']['smallest_area']['ID_name']}) = {diagnostics_obj['source_shape']['smallest_area']['ID_val']}, area = {helper_funcs.format_num(diagnostics_obj['source_shape']['smallest_area']['area'])}, % of total source shape area = {helper_funcs.format_num(diagnostics_obj['source_shape']['smallest_area']['area_percent_total'], digits=5)}",
			'total_area': f"Total source shape area = {helper_funcs.format_num(diagnostics_obj['source_shape']['total_area'])}",
			'avg_area': f"Average source shape area = {helper_funcs.format_num(diagnostics_obj['source_shape']['avg_area'])}",
			'area_smallest_times_avg': f"The smallest source shape area is {helper_funcs.format_num(diagnostics_obj['source_shape']['area_smallest_times_avg'], digits=2)} times smaller than the average",
			'plot': diagnostics_obj['source_shape']['plot'] # maybe dont need the plots
		},
		'target_shape': {
			'title': diagnostics_obj['target_shape']['title'],
			'smallest_area': f"Smallest area (that is also found in the intersection): ID ({diagnostics_obj['target_shape']['smallest_area']['ID_name']}) = {diagnostics_obj['target_shape']['smallest_area']['ID_val']}, area = {helper_funcs.format_num(diagnostics_obj['target_shape']['smallest_area']['area'])}, % of total target shape area = {helper_funcs.format_num(diagnostics_obj['target_shape']['smallest_area']['area_percent_total'], digits=5)}",
			'total_area': f"Total target shape area = {helper_funcs.format_num(diagnostics_obj['target_shape']['total_area'])}",
			'avg_area': f"Average target shape area = {helper_funcs.format_num(diagnostics_obj['target_shape']['avg_area'])}",
			'area_smallest_times_avg': f"The smallest target shape area is {helper_funcs.format_num(diagnostics_obj['target_shape']['area_smallest_times_avg'], digits=2)} times smaller than the average",
			'plot': diagnostics_obj['target_shape']['plot']
		},
		'intersect_shape': {
			'title': diagnostics_obj['intersect_shape']['title'],
			'note': diagnostics_obj['intersect_shape']['note'], 
			'smallest_area': f"Smallest intersected area: ID ({diagnostics_obj['intersect_shape']['smallest_area']['ID_name']}) = {diagnostics_obj['intersect_shape']['smallest_area']['ID_val']}, area = {helper_funcs.format_num(diagnostics_obj['intersect_shape']['smallest_area']['area'])}, % of total intersected shape area = {helper_funcs.format_num(diagnostics_obj['intersect_shape']['smallest_area']['area_percent_total'], digits=5)}",
			'total_area': f"Total target shape area = {helper_funcs.format_num(diagnostics_obj['intersect_shape']['total_area'])}",
			'avg_area': f"Average target shape area = {helper_funcs.format_num(diagnostics_obj['intersect_shape']['avg_area'])}",
			'area_smallest_times_avg': f"The smallest target shape area is {helper_funcs.format_num(diagnostics_obj['intersect_shape']['area_smallest_times_avg'], digits=2)} times smaller than the average",
			'plot': diagnostics_obj['intersect_shape']['plot']
		},
		'tolerance_simulations': {
			'title': 'Tolerance value simulations',
			'simulations': [],
			'plot': diagnostics_obj['tolerance_simulations']['plot']
		}
	}

	sim_obj = diagnostics_obj['tolerance_simulations']['simulations'] # just for shorterning below lines
	for i in range(0, len(sim_obj['tolerance_percent'])):
		sim_obj_curr_text = ("If the tolerance is set at {tolerance_percent}% "
		"({tolerance_units} units), "
		"the total area lost would be {intersect_area_lost} ({intersect_area_lost_percent_total}% of total intersect area), "
		"which costitutes an increase of {intersect_area_lost_change} times "
		"compared to the previous tolerance value, and is the result of removing {intersections_removed} intersections ({intersections_removed_percent}% of total)."
		"The statistics for the new intersected area shape are as follows. "
		"The smallest area\'s ID is {intersect_stats_smallest_area_id} and its area is {intersect_stats_smallest_area} "
		"({intersect_stats_smallest_area_percent_total}% of total intersect area). "
		"The total intersect area is now {intersect_stats_total_area}, and the average intersect area is {intersect_area_avg_area}. "
		"The smallest intersect area is {intersect_stats_area_smallest_times_avg} times smaller than the average intersect area"
		", which constitues a change of {intersect_stats_smallest_times_avg_change} times from the previous tolerance value."
		).format(
				tolerance_percent = helper_funcs.format_num(sim_obj['tolerance_percent'][i], percent=True),
				tolerance_units = helper_funcs.format_num(sim_obj['tolerance_units'][i]),
				intersect_area_lost = helper_funcs.format_num(sim_obj['intersect_area_lost']['area'][i]),
				intersect_area_lost_percent_total = helper_funcs.format_num(sim_obj['intersect_area_lost']['area_percent_total'][i], digits=5),
				intersect_area_lost_change = helper_funcs.format_num(sim_obj['intersect_area_lost']['change'][i], digits=2),
				intersections_removed = helper_funcs.format_num(sim_obj['intersect_area_lost']['intersections_removed'][i]),
				intersections_removed_percent = helper_funcs.format_num(sim_obj['intersect_area_lost']['intersections_removed_percent'][i], digits=2, percent=True),
				intersect_stats_smallest_area_id = sim_obj['intersect_shape_stats']['smallest_area']['ID_val'][i],
				intersect_stats_smallest_area = helper_funcs.format_num(sim_obj['intersect_shape_stats']['smallest_area']['area'][i]),
				intersect_stats_smallest_area_percent_total = helper_funcs.format_num(sim_obj['intersect_shape_stats']['smallest_area']['area_percent_total'][i], digits=5),
				intersect_stats_total_area = helper_funcs.format_num(sim_obj['intersect_shape_stats']['total_area'][i]),
				intersect_area_avg_area = helper_funcs.format_num(sim_obj['intersect_shape_stats']['avg_area'][i]),
				intersect_stats_area_smallest_times_avg = helper_funcs.format_num(sim_obj['intersect_shape_stats']['area_smallest_times_avg'][i], digits=2),
				intersect_stats_smallest_times_avg_change = helper_funcs.format_num(sim_obj['intersect_shape_stats']['smallest_times_avg_change'][i], digits=2)
				)
		diagnostics_obj_text['tolerance_simulations']['simulations'].append(sim_obj_curr_text)

	return diagnostics_obj_text

def get_tolerance_simulations_obj(tolerance_values_sim_prop, intersect_unedited):
	tolerance_simulations_obj = {
		'tolerance_percent': [],
		'tolerance_units': [],
		'intersect_area_lost': {
			'area': [],
			'area_percent_total': [],
			'change': [],
			'intersections_removed': [],
			'intersections_removed_percent': []
		},
		'intersect_shape_stats': {
			'smallest_area': {
				'ID_val': [],
				'area': [],
				'area_percent_total': [],
			},
			'total_area': [],
			'avg_area': [],
			'area_smallest_times_avg': [],
			'smallest_times_avg_change': []
		}
	}

	for i, tolerance_prop in enumerate(tolerance_values_sim_prop):
		tolerance_simulations_obj['tolerance_percent'].append(tolerance_prop*100)
		intersect_smallest_area_postfilter = min(min(intersect_unedited['area_base_source']), min(intersect_unedited['area_base_target']))*tolerance_prop
		tolerance_simulations_obj['tolerance_units'].append(intersect_smallest_area_postfilter)
		
		# note the sign here - we only keep what is to be later removed
		intersect_area_lost = intersect_unedited[intersect_unedited['intersect_area'] <= intersect_smallest_area_postfilter]['intersect_area'].sum()
		tolerance_simulations_obj['intersect_area_lost']['area'].append(intersect_area_lost)
		tolerance_simulations_obj['intersect_area_lost']['area_percent_total'].append(intersect_area_lost / intersect_unedited['intersect_area'].sum())
		intersections_removed_temp = sum(intersect_unedited['intersect_area'] < intersect_smallest_area_postfilter)
		tolerance_simulations_obj['intersect_area_lost']['intersections_removed'].append(intersections_removed_temp)
		tolerance_simulations_obj['intersect_area_lost']['intersections_removed_percent'].append(round(intersections_removed_temp/intersect_unedited.shape[0]*100, 2))

		intersect_filtered = intersect_unedited[intersect_unedited['intersect_area'] > intersect_smallest_area_postfilter]

		if len(intersect_filtered) < 1:
			# in case the tolerance eliminates all areas
			tolerance_simulations_obj['intersect_shape_stats']['smallest_area']['ID_val'].append('NA')
			tolerance_simulations_obj['intersect_shape_stats']['smallest_area']['area'].append(np.nan)
			tolerance_simulations_obj['intersect_shape_stats']['smallest_area']['area_percent_total'].append(np.nan)
			tolerance_simulations_obj['intersect_shape_stats']['total_area'].append(np.nan)
			tolerance_simulations_obj['intersect_shape_stats']['avg_area'].append(np.nan)
			tolerance_simulations_obj['intersect_shape_stats']['area_smallest_times_avg'].append(np.nan)
		else:
			tolerance_simulations_obj['intersect_shape_stats']['smallest_area']['ID_val'].append(intersect_filtered.sort_values('intersect_area')['INTERSECT_ID'].iloc[0])
			tolerance_simulations_obj['intersect_shape_stats']['smallest_area']['area'].append(min(intersect_filtered['intersect_area']))
			tolerance_simulations_obj['intersect_shape_stats']['smallest_area']['area_percent_total'].append(min(intersect_filtered['intersect_area']) / sum(intersect_filtered['intersect_area']))
			tolerance_simulations_obj['intersect_shape_stats']['total_area'].append(sum(intersect_filtered['intersect_area']))
			tolerance_simulations_obj['intersect_shape_stats']['avg_area'].append(np.mean(intersect_filtered['intersect_area']))
			tolerance_simulations_obj['intersect_shape_stats']['area_smallest_times_avg'].append(np.mean(intersect_filtered['intersect_area']) / min(intersect_filtered['intersect_area']))


		if i == 0:
			intersect_area_lost_change = 0
			intersect_smallest_times_avg_change = 0
		else:
			with np.errstate(divide='ignore', invalid='ignore'):
				if len(intersect_filtered) < 1:
					# in case the tolerance eliminates all areas
					intersect_area_lost_change = np.nan
					intersect_smallest_times_avg_change = np.nan
				else:
					intersect_area_lost_change = tolerance_simulations_obj['intersect_area_lost']['area'][i] / tolerance_simulations_obj['intersect_area_lost']['area'][i-1]
					if np.isnan(intersect_area_lost_change) or np.isinf(intersect_area_lost_change):
						intersect_area_lost_change = 0

					intersect_smallest_times_avg_change = tolerance_simulations_obj['intersect_shape_stats']['area_smallest_times_avg'][i] / tolerance_simulations_obj['intersect_shape_stats']['area_smallest_times_avg'][i-1]
					if np.isnan(intersect_smallest_times_avg_change) or np.isinf(intersect_smallest_times_avg_change):
						intersect_smallest_times_avg_change = 0

		tolerance_simulations_obj['intersect_area_lost']['change'].append(intersect_area_lost_change)
		tolerance_simulations_obj['intersect_shape_stats']['smallest_times_avg_change'].append(intersect_smallest_times_avg_change)

	return tolerance_simulations_obj


def save_diagnostics_to_word(diagnostics_obj, output_type, filename = None):
	doc = Document()
	diagnostics_obj_text = get_diagnostics_obj_text(diagnostics_obj)

	if output_type == 'text':
		doc.add_heading(diagnostics_obj_text['title'], 0)
		doc.add_paragraph(diagnostics_obj_text['units'])
		doc.add_paragraph(diagnostics_obj_text['area_lost_during_intersection'])
		doc.add_heading('1. Statistics about shapes', level=1)
		doc.add_heading('1.1 ' + diagnostics_obj_text['source_shape']['title'], level=2)
		doc.add_paragraph(diagnostics_obj_text['source_shape']['smallest_area'])
		doc.add_paragraph(diagnostics_obj_text['source_shape']['total_area'])
		doc.add_paragraph(diagnostics_obj_text['source_shape']['avg_area'])
		doc.add_paragraph(diagnostics_obj_text['source_shape']['area_smallest_times_avg'])
		doc.add_heading('1.2 ' + diagnostics_obj_text['target_shape']['title'], level=2)
		doc.add_paragraph(diagnostics_obj_text['target_shape']['smallest_area'])
		doc.add_paragraph(diagnostics_obj_text['target_shape']['total_area'])
		doc.add_paragraph(diagnostics_obj_text['target_shape']['avg_area'])
		doc.add_paragraph(diagnostics_obj_text['target_shape']['area_smallest_times_avg'])
		doc.add_heading('1.3 ' + diagnostics_obj_text['intersect_shape']['title'], level=2)
		doc.add_paragraph(diagnostics_obj_text['intersect_shape']['note'])
		doc.add_paragraph(diagnostics_obj_text['intersect_shape']['smallest_area'])
		doc.add_paragraph(diagnostics_obj_text['intersect_shape']['total_area'])
		doc.add_paragraph(diagnostics_obj_text['intersect_shape']['avg_area'])
		doc.add_paragraph(diagnostics_obj_text['intersect_shape']['area_smallest_times_avg'])
		doc.add_heading('2. ' + diagnostics_obj_text['tolerance_simulations']['title'], level=1)
		for sim in diagnostics_obj_text['tolerance_simulations']['simulations']:
			doc.add_paragraph(sim)

		doc.add_heading('3. Source and target shapes visualisation', level=1)

		doc.add_picture(diagnostics_obj_text['source_shape']['plot'])
		doc.add_picture(diagnostics_obj_text['target_shape']['plot'])
		doc.add_picture(diagnostics_obj_text['intersect_shape']['plot'])

	elif output_type == 'table':
		helper_funcs.change_orinetation(doc)
		doc.add_heading(diagnostics_obj['title'], 0)
		doc.add_paragraph(diagnostics_obj_text['units'])
		doc.add_paragraph(diagnostics_obj_text['area_lost_during_intersection'])
		doc.add_heading('1. Statistics about shapes', level=1)
		
		shape_stats_table = doc.add_table(rows=5, cols=8)
		shape_stats_table.style = 'Table Grid'

		# merge smallest area
		shape_stats_table.cell(row_idx=0, col_idx=1).merge(shape_stats_table.cell(row_idx=0, col_idx=4))
		# merge other 3 cols
		shape_stats_table.cell(row_idx=0, col_idx=5).merge(shape_stats_table.cell(row_idx=1, col_idx=5))
		shape_stats_table.cell(row_idx=0, col_idx=6).merge(shape_stats_table.cell(row_idx=1, col_idx=6))
		shape_stats_table.cell(row_idx=0, col_idx=7).merge(shape_stats_table.cell(row_idx=1, col_idx=7))

		shape_stats_table.cell(row_idx=0, col_idx=1).text = 'Smallest area'

		shape_stats_table.cell(row_idx=1, col_idx=1).text = 'ID column name'
		shape_stats_table.cell(row_idx=1, col_idx=2).text = 'ID'
		shape_stats_table.cell(row_idx=1, col_idx=3).text = 'Area'
		shape_stats_table.cell(row_idx=1, col_idx=4).text = '% total shape area'

		shape_stats_table.cell(row_idx=0, col_idx=5).text = 'Total shape area'
		shape_stats_table.cell(row_idx=0, col_idx=6).text = 'Average shape area'
		shape_stats_table.cell(row_idx=0, col_idx=7).text = 'Times the smallest shape area smaller than the average'

		shape_stats_table.cell(row_idx=2, col_idx=0).text = 'Source shape'
		shape_stats_table.cell(row_idx=3, col_idx=0).text = 'Target shape'
		shape_stats_table.cell(row_idx=4, col_idx=0).text = 'Intersected shape*'

		for row_idx in range(2, 5):
			diagnostics_obj_section_lookup = {2: 'source_shape', 3: 'target_shape', 4: 'intersect_shape'}
			shape_stats_table.cell(row_idx=row_idx, col_idx=1).text = diagnostics_obj[diagnostics_obj_section_lookup[row_idx]]['smallest_area']['ID_name']
			shape_stats_table.cell(row_idx=row_idx, col_idx=2).text = diagnostics_obj[diagnostics_obj_section_lookup[row_idx]]['smallest_area']['ID_val']
			shape_stats_table.cell(row_idx=row_idx, col_idx=3).text = helper_funcs.format_num(diagnostics_obj[diagnostics_obj_section_lookup[row_idx]]['smallest_area']['area'])
			shape_stats_table.cell(row_idx=row_idx, col_idx=4).text = helper_funcs.format_num(diagnostics_obj[diagnostics_obj_section_lookup[row_idx]]['smallest_area']['area_percent_total'], digits=5)
			shape_stats_table.cell(row_idx=row_idx, col_idx=5).text = helper_funcs.format_num(diagnostics_obj[diagnostics_obj_section_lookup[row_idx]]['total_area'])
			shape_stats_table.cell(row_idx=row_idx, col_idx=6).text = helper_funcs.format_num(diagnostics_obj[diagnostics_obj_section_lookup[row_idx]]['avg_area'])
			shape_stats_table.cell(row_idx=row_idx, col_idx=7).text = helper_funcs.format_num(diagnostics_obj[diagnostics_obj_section_lookup[row_idx]]['area_smallest_times_avg'], digits=2)

		shape_stats_note = doc.add_paragraph()
		shape_stats_note.add_run("* Note: ").bold = True
		shape_stats_note.add_run(diagnostics_obj['intersect_shape']['note'])
		
		doc.add_heading('2. Tolerance value simulations', level=1)
		sim_obj = diagnostics_obj['tolerance_simulations']['simulations'] # just for shorterning below code

		doc.add_heading('2.1. Area lost statistics', level=2)

		tolerance_area_lost_table = doc.add_table(rows=1+len(sim_obj['tolerance_percent']), cols=7)
		tolerance_area_lost_table.style = 'Table Grid'

		tolerance_area_lost_table.cell(row_idx=0, col_idx=0).text = 'Tolerance percent'
		tolerance_area_lost_table.cell(row_idx=0, col_idx=1).text = f"Tolerance units (in {diagnostics_obj['units']})"
		tolerance_area_lost_table.cell(row_idx=0, col_idx=2).text = 'Intersected area lost'
		tolerance_area_lost_table.cell(row_idx=0, col_idx=3).text = '% lost of total intersected area'
		tolerance_area_lost_table.cell(row_idx=0, col_idx=4).text = 'Times increase from previous tolerance threshold'
		tolerance_area_lost_table.cell(row_idx=0, col_idx=5).text = 'Number of intersections removed'
		tolerance_area_lost_table.cell(row_idx=0, col_idx=6).text = '% of intersections removed'


		for i in range(0, len(sim_obj['tolerance_percent'])):
			tolerance_area_lost_table.cell(row_idx=i+1, col_idx=0).text = helper_funcs.format_num(sim_obj['tolerance_percent'][i], percent=True)
			tolerance_area_lost_table.cell(row_idx=i+1, col_idx=1).text = helper_funcs.format_num(sim_obj['tolerance_units'][i])
			tolerance_area_lost_table.cell(row_idx=i+1, col_idx=2).text = helper_funcs.format_num(sim_obj['intersect_area_lost']['area'][i])
			tolerance_area_lost_table.cell(row_idx=i+1, col_idx=3).text = helper_funcs.format_num(sim_obj['intersect_area_lost']['area_percent_total'][i], digits=5)
			tolerance_area_lost_table.cell(row_idx=i+1, col_idx=4).text = helper_funcs.format_num(sim_obj['intersect_area_lost']['change'][i], digits=2)
			tolerance_area_lost_table.cell(row_idx=i+1, col_idx=5).text = helper_funcs.format_num(sim_obj['intersect_area_lost']['intersections_removed'][i], digits=2)
			tolerance_area_lost_table.cell(row_idx=i+1, col_idx=6).text = helper_funcs.format_num(sim_obj['intersect_area_lost']['intersections_removed_percent'][i], digits=2, percent=True)

		doc.add_heading('2.2. New intersected shape statistics', level=2)

		tolerance_intersect_stats_table = doc.add_table(rows=2+len(sim_obj['tolerance_percent']), cols=9)
		tolerance_intersect_stats_table.style = 'Table Grid'

		# merge smallest area
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=2).merge(tolerance_intersect_stats_table.cell(row_idx=0, col_idx=4))
		# merge header cols
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=0).merge(tolerance_intersect_stats_table.cell(row_idx=1, col_idx=0))
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=1).merge(tolerance_intersect_stats_table.cell(row_idx=1, col_idx=1))
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=5).merge(tolerance_intersect_stats_table.cell(row_idx=1, col_idx=5))
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=6).merge(tolerance_intersect_stats_table.cell(row_idx=1, col_idx=6))
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=7).merge(tolerance_intersect_stats_table.cell(row_idx=1, col_idx=7))
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=8).merge(tolerance_intersect_stats_table.cell(row_idx=1, col_idx=8))

		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=0).text = 'Tolerance percent'
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=1).text = f"Tolerance units (in {diagnostics_obj['units']})"
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=2).text = 'Smallest area'
		tolerance_intersect_stats_table.cell(row_idx=1, col_idx=2).text = 'ID'
		tolerance_intersect_stats_table.cell(row_idx=1, col_idx=3).text = 'Area'
		tolerance_intersect_stats_table.cell(row_idx=1, col_idx=4).text = '% total shape area'
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=5).text = 'Total shape area'
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=6).text = 'Average shape area'
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=7).text = 'Times the smallest shape area smaller than the average'
		tolerance_intersect_stats_table.cell(row_idx=0, col_idx=8).text = 'Times change in the smallest to average shape area'

		for i in range(0, len(sim_obj['tolerance_percent'])):
			tolerance_intersect_stats_table.cell(row_idx=i+2, col_idx=0).text = helper_funcs.format_num(sim_obj['tolerance_percent'][i], percent=True)
			tolerance_intersect_stats_table.cell(row_idx=i+2, col_idx=1).text = helper_funcs.format_num(sim_obj['tolerance_units'][i])
			tolerance_intersect_stats_table.cell(row_idx=i+2, col_idx=2).text = sim_obj['intersect_shape_stats']['smallest_area']['ID_val'][i]
			tolerance_intersect_stats_table.cell(row_idx=i+2, col_idx=3).text = helper_funcs.format_num(sim_obj['intersect_shape_stats']['smallest_area']['area'][i])
			tolerance_intersect_stats_table.cell(row_idx=i+2, col_idx=4).text = helper_funcs.format_num(sim_obj['intersect_shape_stats']['smallest_area']['area_percent_total'][i], digits=5)
			tolerance_intersect_stats_table.cell(row_idx=i+2, col_idx=5).text = helper_funcs.format_num(sim_obj['intersect_shape_stats']['total_area'][i])
			tolerance_intersect_stats_table.cell(row_idx=i+2, col_idx=6).text = helper_funcs.format_num(sim_obj['intersect_shape_stats']['avg_area'][i])
			tolerance_intersect_stats_table.cell(row_idx=i+2, col_idx=7).text = helper_funcs.format_num(sim_obj['intersect_shape_stats']['area_smallest_times_avg'][i], digits=2)
			tolerance_intersect_stats_table.cell(row_idx=i+2, col_idx=8).text = helper_funcs.format_num(sim_obj['intersect_shape_stats']['smallest_times_avg_change'][i], digits=2)

		doc.add_heading('3. Source and target shapes visualisation', level=2)			

		doc.add_picture(diagnostics_obj_text['source_shape']['plot'])
		doc.add_picture(diagnostics_obj_text['target_shape']['plot'])
		doc.add_picture(diagnostics_obj_text['intersect_shape']['plot'])

		doc = helper_funcs.set_autofit(doc)

	if filename is None:
		time_now = datetime.now().strftime('%H%M%S-%Y%m%d')
		filename = 'crosswalk_diagnostics_'+time_now+'.docx'
	if not(filename.endswith('.docx')):
		filename = filename+'.docx'

	doc.save(filename)
	

def print_diagnostics(diagnostics_obj):
	diagnostics_obj_text = get_diagnostics_obj_text(diagnostics_obj)

	diagnostics_text_arr = []
	diagnostics_text_arr.append(f"--------------------------------{diagnostics_obj_text['title']}------------------------------")
	diagnostics_text_arr.append(diagnostics_obj_text['units'])
	diagnostics_text_arr.append(diagnostics_obj_text['area_lost_during_intersection'])
	diagnostics_text_arr.append(f"--------------{diagnostics_obj_text['source_shape']['title']}--------------------")
	diagnostics_text_arr.append(diagnostics_obj_text['source_shape']['smallest_area'])
	diagnostics_text_arr.append(diagnostics_obj_text['source_shape']['total_area'])
	diagnostics_text_arr.append(diagnostics_obj_text['source_shape']['avg_area'])
	diagnostics_text_arr.append(diagnostics_obj_text['source_shape']['area_smallest_times_avg'])
	diagnostics_text_arr.append(f"--------------{diagnostics_obj_text['target_shape']['title']}--------------------")
	diagnostics_text_arr.append(diagnostics_obj_text['target_shape']['smallest_area'])
	diagnostics_text_arr.append(diagnostics_obj_text['target_shape']['total_area'])
	diagnostics_text_arr.append(diagnostics_obj_text['target_shape']['avg_area'])
	diagnostics_text_arr.append(diagnostics_obj_text['target_shape']['area_smallest_times_avg'])
	diagnostics_text_arr.append(f"--------------{diagnostics_obj_text['intersect_shape']['title']}--------------------")
	diagnostics_text_arr.append(diagnostics_obj_text['intersect_shape']['note'])
	diagnostics_text_arr.append(diagnostics_obj_text['intersect_shape']['smallest_area'])
	diagnostics_text_arr.append(diagnostics_obj_text['intersect_shape']['total_area'])
	diagnostics_text_arr.append(diagnostics_obj_text['intersect_shape']['avg_area'])
	diagnostics_text_arr.append(diagnostics_obj_text['intersect_shape']['area_smallest_times_avg'])
	diagnostics_text_arr.append(f"--------------{diagnostics_obj_text['tolerance_simulations']['title']}--------------------")
	diagnostics_text_arr = diagnostics_text_arr + diagnostics_obj_text['tolerance_simulations']['simulations']

	print('\n'.join(diagnostics_text_arr))