from docx import Document
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from docx.table import _Cell
# from docx.shared import Pt
# from docx import Document
# from docx.shared import Inches
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT

def get_projection_units(source_shape, target_shape):
	source_axes = source_shape.crs.axis_info
	target_axes = target_shape.crs.axis_info
	if (source_axes[0].unit_name != source_axes[1].unit_name) or (target_axes[0].unit_name != target_axes[0].unit_name):
		raise ProjectionException('Source or target axes\' projection units are different.')

	if source_axes[0].unit_name != target_axes[0].unit_name:
		raise ProjectionException('Source projection units are different from target projection units.')

	return source_axes[0].unit_name

def format_num(x, digits=0, percent=False):
	if x < 0:
		raise Exception("Attempt to format negative number - this is currently not supported.")
	if digits == 'first_non_zero':
		# finds count of digits after decimal point until first non-zero
		digits = int(('%e' % x).partition('-')[2])
	if percent and digits == 0:
		if x < 1:
			digits = 3
		else:
			digits = 0

	res = '{:.{}f}'.format(x, digits)
	return res

def set_autofit(doc: Document) -> Document:
	#credits - https://github.com/python-openxml/python-docx/issues/209#issuecomment-566128709
	"""
	Hotfix for autofit.
	"""
	for t_idx, table in enumerate(doc.tables):
		doc.tables[t_idx].autofit = True
		doc.tables[t_idx].allow_autofit = True
		doc.tables[t_idx]._tblPr.xpath("./w:tblW")[0].attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
		for row_idx, r_val in enumerate(doc.tables[t_idx].rows):
			for cell_idx, c_val in enumerate(doc.tables[t_idx].rows[row_idx].cells):
				doc.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.type = 'auto'
				doc.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.w = 0
	return doc

def change_orinetation(document):
	# not among the most flexible of functionalities; be careful, see https://python-docx.readthedocs.io/en/latest/user/sections.html#working-with-sections
    current_section = document.sections[0]
    new_width, new_height = current_section.page_height, current_section.page_width
    current_section.orientation = WD_ORIENT.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height

    return current_section
